from __future__ import annotations

"""Template-based document generation engine for the IQW platform.

Provides ORM-backed template storage, Jinja2-powered placeholder substitution,
and export to DOCX / PDF formats.  All public functions return plain dicts so
the API layer can serialise them directly.
"""

import io
import re
import uuid as _uuid
from datetime import datetime, timezone
from typing import Dict, List, Optional

import sqlalchemy as sa
from jinja2 import Template
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from audit import compute_sha256
from models import DocumentTemplate, GeneratedDocument

# ---------------------------------------------------------------------------
# DEPRECATED: In-memory stores kept for backward compatibility only.
# These dicts are no longer populated by the module — all reads and writes
# now go through the async SQLAlchemy ORM layer.  Do NOT rely on them in
# new code; they will be removed in a future release.
# ---------------------------------------------------------------------------

_templates: Dict[str, dict] = {}
_generated_documents: Dict[str, dict] = {}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def compute_generated_document_hash(content: Optional[str]) -> str:
    """Return the SHA-256 hash for generated document text content."""
    return compute_sha256((content or "").encode("utf-8"))


def _refresh_generated_document_hash(doc: GeneratedDocument) -> None:
    doc.sha256_hash = compute_generated_document_hash(doc.generated_content)


def extract_placeholders(template_body: str) -> List[str]:
    """Return a deduplicated, order-preserving list of ``{{...}}`` tokens."""
    seen: set = set()
    result: List[str] = []
    for match in _PLACEHOLDER_RE.finditer(template_body):
        name = match.group(1)
        if name not in seen:
            seen.add(name)
            result.append(name)
    return result


# ---------------------------------------------------------------------------
# ORM → dict helpers
# ---------------------------------------------------------------------------

def _template_to_dict(tpl: DocumentTemplate) -> dict:
    return {
        "id": tpl.id,
        "template_name": tpl.template_name,
        "category": tpl.category.value if hasattr(tpl.category, "value") else tpl.category,
        "template_body": tpl.template_body,
        "placeholders": tpl.placeholders or [],
        "is_active": tpl.is_active,
        "version": tpl.version,
    }


def _gendoc_to_dict(doc: GeneratedDocument, auto_filled: list, missing: list) -> dict:
    signature_status = doc.digital_signature_status.value if hasattr(doc.digital_signature_status, "value") else doc.digital_signature_status
    return {
        "id": doc.id,
        "template_id": doc.template_id,
        "case_id": doc.case_id,
        "category": doc.document_category.value if hasattr(doc.document_category, "value") else doc.document_category,
        "content": doc.generated_content,
        "sha256": doc.sha256_hash,
        "auto_filled_fields": auto_filled,
        "missing_fields": missing,
        "missing_prompt": (
            "Please complete the following case details before generating this document: "
            + ", ".join(missing)
            + "."
            if missing else None
        ),
        "digital_signature_status": signature_status,
        "is_read_only": signature_status == "Signed",
        "signed_by": doc.signed_by,
        "signed_at": doc.signed_at.isoformat() if doc.signed_at else None,
        "signature_certificate_details": doc.signature_certificate_details,
        "created_by": doc.created_by,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
        "io_edited": doc.io_edited,
        "updated_by": doc.updated_by,
    }


# ---------------------------------------------------------------------------
# Template seeding
# ---------------------------------------------------------------------------

async def seed_templates(db: AsyncSession) -> None:
    """Populate ``document_templates`` with the 13 standard IQW document templates."""

    # Check if templates already seeded
    result = await db.execute(select(sa.func.count()).select_from(DocumentTemplate))
    count = result.scalar() or 0
    if count > 0:
        return  # already seeded

    defs: List[dict] = [
        # ---- FSL Communications (3) ----
        {
            "id": "tpl-fsl-001",
            "template_name": "FSL Forwarding Letter",
            "category": "FSL_Communication",
            "template_body": (
                "To,\nThe Director,\n{{fsl_lab_name}}\n\n"
                "Subject: Forwarding of case property for forensic examination\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n\n"
                "Sir/Madam,\n\n"
                "It is respectfully submitted that during the course of investigation "
                "of the above-referred case, the following articles/exhibits have been "
                "seized and are being forwarded herewith for forensic examination:\n\n"
                "Description of Evidence: {{evidence_description}}\n\n"
                "The seized articles are packed, sealed and labelled in the presence "
                "of independent witnesses as per prescribed procedure.\n\n"
                "You are requested to kindly examine the said articles and furnish "
                "the forensic analysis report at the earliest.\n\n"
                "Accused Name: {{accused_name}}\n\n"
                "Yours faithfully,\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-fsl-002",
            "template_name": "Sample Forwarding Memo",
            "category": "FSL_Communication",
            "template_body": (
                "SAMPLE FORWARDING MEMO\n\n"
                "Crime No.: {{case_number}}    P.S.: {{police_station}}\n"
                "Date: {{date}}\n\n"
                "The following samples are forwarded to the Forensic Science "
                "Laboratory for examination:\n\n"
                "Sample Description: {{sample_description}}\n"
                "Number of Samples: {{sample_count}}\n"
                "Sealed With: {{seal_details}}\n\n"
                "Purpose of Examination: {{examination_purpose}}\n\n"
                "Forwarded by: {{io_name}}, {{io_rank}}\n"
                "P.S. {{police_station}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-fsl-003",
            "template_name": "FSL Reminder Letter",
            "category": "FSL_Communication",
            "template_body": (
                "To,\nThe Director,\n{{fsl_lab_name}}\n\n"
                "Subject: Reminder for pending FSL report\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n"
                "Original forwarding date: {{original_forwarding_date}}\n\n"
                "Sir/Madam,\n\n"
                "This is to bring to your kind attention that the forensic "
                "examination report pertaining to the above case is still "
                "awaited. The exhibits were forwarded vide this office memo "
                "dated {{original_forwarding_date}}.\n\n"
                "The report is urgently required for the purpose of "
                "investigation and filing of the charge sheet within the "
                "statutory time limit.\n\n"
                "You are requested to expedite the examination and furnish "
                "the report at the earliest.\n\n"
                "Yours faithfully,\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },

        # ---- Evidence Certificates (2) ----
        {
            "id": "tpl-evd-001",
            "template_name": "Section 63 BSA Certificate",
            "category": "Evidence_Certificate",
            "template_body": (
                "CERTIFICATE UNDER SECTION 63 OF THE "
                "BHARATIYA SAKSHYA ADHINIYAM, 2023\n\n"
                "Crime No.: {{case_number}}\n"
                "Date: {{date}}\n\n"
                "I, {{io_name}}, hereby certify that:\n\n"
                "1. The electronic record described below was produced by a "
                "computer during the period in which the computer was used "
                "regularly to store or process information.\n\n"
                "2. During the said period, information of the kind contained "
                "in the electronic record was regularly fed into the computer "
                "in the ordinary course of activities.\n\n"
                "3. The computer was operating properly throughout the material "
                "part of the said period, and if not, any malfunction did not "
                "affect the electronic record or its accuracy.\n\n"
                "4. The information contained in the electronic record "
                "reproduces or is derived from information fed into the "
                "computer in the ordinary course of activities.\n\n"
                "Document Description: {{document_description}}\n"
                "Hash Value (SHA-256): {{hash_value}}\n\n"
                "Signature: ____________________\n"
                "{{io_name}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-evd-002",
            "template_name": "Hash Value Declaration",
            "category": "Evidence_Certificate",
            "template_body": (
                "HASH VALUE DECLARATION\n\n"
                "Crime No.: {{case_number}}\n"
                "Date: {{date}}\n\n"
                "I, {{io_name}}, the Investigating Officer in the above case, "
                "do hereby declare that the following hash values were computed "
                "at the time of seizure of the digital evidence:\n\n"
                "Evidence Description: {{evidence_description}}\n"
                "Algorithm: SHA-256\n"
                "Hash Value: {{hash_value}}\n\n"
                "The hash value was computed using a forensically sound tool "
                "and verified in the presence of independent witnesses.\n\n"
                "Witness 1: {{witness_1}}\n"
                "Witness 2: {{witness_2}}\n\n"
                "Signature: ____________________\n"
                "{{io_name}}"
            ),
            "version": 1,
        },

        # ---- Legal Notices (4) ----
        {
            "id": "tpl-lgn-001",
            "template_name": "Bank Account Information Request",
            "category": "Legal_Notice",
            "template_body": (
                "To,\nThe Branch Manager,\n{{bank_name}}, {{branch_name}}\n\n"
                "Subject: Request for bank account details under Section 94 "
                "BNSS / Section 91 CrPC\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n\n"
                "Sir/Madam,\n\n"
                "During the investigation of the above case, it has become "
                "necessary to obtain the bank account details of the following "
                "person:\n\n"
                "Account Holder: {{account_holder_name}}\n"
                "Account Number: {{account_number}}\n\n"
                "You are hereby requested to furnish the following information "
                "within 72 hours:\n"
                "1. Complete KYC documents\n"
                "2. Account opening form\n"
                "3. Statement of account for the last 12 months\n"
                "4. Details of linked accounts and beneficiaries\n\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgn-002",
            "template_name": "Bank Account Freeze Request",
            "category": "Legal_Notice",
            "template_body": (
                "URGENT\n\n"
                "To,\nThe Branch Manager,\n{{bank_name}}, {{branch_name}}\n\n"
                "Subject: Request to freeze bank account\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n\n"
                "Sir/Madam,\n\n"
                "In connection with the investigation of the above-mentioned "
                "case, there are reasonable grounds to believe that the "
                "following bank account contains proceeds of crime:\n\n"
                "Account Holder: {{account_holder_name}}\n"
                "Account Number: {{account_number}}\n\n"
                "You are hereby requested to immediately freeze all debits "
                "from the above account pending further orders. Credits may "
                "continue to be allowed.\n\n"
                "A formal court order will follow within the prescribed "
                "statutory period.\n\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgn-003",
            "template_name": "ISP Data Request",
            "category": "Legal_Notice",
            "template_body": (
                "To,\nThe Nodal Officer (Law Enforcement),\n{{isp_name}}\n\n"
                "Subject: Request for subscriber and traffic data under "
                "Section 94 BNSS\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n\n"
                "Sir/Madam,\n\n"
                "The following information is required in connection with "
                "the investigation of a cognizable offence:\n\n"
                "IP Address: {{ip_address}}\n"
                "Date/Time Range: {{date_range}}\n"
                "Email ID (if applicable): {{email_id}}\n\n"
                "Please furnish:\n"
                "1. Subscriber registration details including KYC\n"
                "2. Login/session logs for the specified period\n"
                "3. Associated MAC addresses and device identifiers\n\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgn-004",
            "template_name": "CDR Requisition",
            "category": "Legal_Notice",
            "template_body": (
                "To,\nThe Nodal Officer (Law Enforcement),\n"
                "{{telecom_provider}}\n\n"
                "Subject: Requisition for Call Detail Records under "
                "Section 94 BNSS\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n\n"
                "Sir/Madam,\n\n"
                "In connection with the investigation of the above case, "
                "you are requested to furnish the following:\n\n"
                "Mobile Number: {{mobile_number}}\n"
                "IMEI Number (if known): {{imei_number}}\n"
                "Period: {{cdr_period}}\n\n"
                "Data Required:\n"
                "1. Incoming and outgoing call records\n"
                "2. SMS details (incoming and outgoing)\n"
                "3. Cell tower location data (Cell ID)\n"
                "4. Recharge history and data usage logs\n"
                "5. Subscriber details and KYC documents\n\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgn-005",
            "template_name": "Google Platform Data Request",
            "category": "Legal_Notice",
            "template_body": (
                "To,\nThe Law Enforcement Response Team,\nGoogle LLC\n\n"
                "Subject: Platform data preservation and disclosure request\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n\n"
                "Please preserve and provide subscriber, login, IP, and content metadata "
                "available for the following identifier in connection with the investigation:\n\n"
                "Google Account / Email / Channel ID: {{platform_identifier}}\n"
                "Date Range: {{date_range}}\n"
                "Legal Basis: {{legal_basis}}\n\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgn-006",
            "template_name": "Meta Platform Data Request",
            "category": "Legal_Notice",
            "template_body": (
                "To,\nThe Law Enforcement Response Team,\nMeta Platforms, Inc.\n\n"
                "Subject: Preservation and disclosure request for platform records\n\n"
                "Ref: Crime No. {{case_number}}, P.S. {{police_station}}\n\n"
                "Please preserve and disclose available account, login, IP, and message "
                "metadata for the following identifier:\n\n"
                "Facebook / Instagram / WhatsApp Identifier: {{platform_identifier}}\n"
                "Date Range: {{date_range}}\n"
                "Legal Basis: {{legal_basis}}\n\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },

        # ---- Legal Drafts (4) ----
        {
            "id": "tpl-lgd-001",
            "template_name": "Arrest Memo",
            "category": "Legal_Draft",
            "template_body": (
                "ARREST MEMO\n"
                "(Under Section 35(4) of BNSS, 2023)\n\n"
                "Crime No.: {{case_number}}\n\n"
                "1. Name of Arrested Person: {{accused_name}}\n"
                "2. Date and Time of Arrest: {{date_time_arrest}}\n"
                "3. Place of Arrest: {{place_arrest}}\n\n"
                "4. The arrested person has been informed of:\n"
                "   (a) The grounds of arrest\n"
                "   (b) The right to have a person informed of the arrest\n"
                "   (c) The right to consult a legal practitioner of choice\n\n"
                "5. Arresting Officer: {{io_name}}\n\n"
                "6. Witnesses to arrest:\n{{witnesses}}\n\n"
                "7. The arrested person was medically examined at the time "
                "of arrest and found to be in the following condition: "
                "{{medical_condition}}\n\n"
                "Signature of Arrested Person: ____________________\n"
                "Signature of IO: ____________________\n"
                "Signature of Witnesses: ____________________"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgd-002",
            "template_name": "Seizure Memo",
            "category": "Legal_Draft",
            "template_body": (
                "SEIZURE MEMO / PANCHNAMA\n"
                "(Under Section 105 of BNSS, 2023)\n\n"
                "Crime No.: {{case_number}}\n"
                "Date: {{date}}\n"
                "Place of Seizure: {{place_of_seizure}}\n\n"
                "I, {{io_name}}, {{io_rank}}, the Investigating Officer, "
                "in the presence of the following witnesses, do hereby seize "
                "the following articles:\n\n"
                "{{seized_items}}\n\n"
                "The above articles have been packed, sealed and labelled "
                "with the following seal impression: {{seal_impression}}\n\n"
                "Witness 1: {{witness_1}}\n"
                "Witness 2: {{witness_2}}\n\n"
                "Signature of IO: ____________________\n"
                "Signature of Witnesses: ____________________"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgd-003",
            "template_name": "Remand Note",
            "category": "Legal_Draft",
            "template_body": (
                "IN THE COURT OF {{court_name}}\n\n"
                "Crime No.: {{case_number}}\n"
                "P.S.: {{police_station}}\n\n"
                "APPLICATION FOR POLICE/JUDICIAL CUSTODY REMAND\n"
                "(Under Section 187 of BNSS, 2023)\n\n"
                "Respectfully Showeth:\n\n"
                "1. The accused {{accused_name}} was arrested on "
                "{{date_of_arrest}} in connection with the above crime.\n\n"
                "2. Brief facts of the case:\n{{brief_facts}}\n\n"
                "3. The following investigation steps are pending:\n"
                "{{pending_investigation}}\n\n"
                "4. It is prayed that the accused may kindly be remanded "
                "to police/judicial custody for a period of {{remand_days}} "
                "days for the purpose of further investigation.\n\n"
                "{{io_name}}, {{io_rank}}\n"
                "Date: {{date}}"
            ),
            "version": 1,
        },
        {
            "id": "tpl-lgd-004",
            "template_name": "Confession Recording Template",
            "category": "Legal_Draft",
            "template_body": (
                "RECORD OF CONFESSION\n"
                "(Under Section 183 of BNSS, 2023 read with Section 23 "
                "of BSA, 2023)\n\n"
                "Crime No.: {{case_number}}\n"
                "Date: {{date}}\n\n"
                "Before: {{magistrate_name}}, {{magistrate_designation}}\n\n"
                "Name of Accused: {{accused_name}}\n"
                "Age: {{accused_age}}\n"
                "Address: {{accused_address}}\n\n"
                "The accused was produced before me on {{date}}. "
                "After being warned that:\n"
                "(a) He/she is not bound to make a confession;\n"
                "(b) Any confession made may be used as evidence against "
                "him/her;\n"
                "(c) He/she has been given adequate time for reflection;\n\n"
                "The accused voluntarily states as follows:\n\n"
                "{{confession_text}}\n\n"
                "The above statement was recorded in my presence and "
                "the accused has signed/thumb-impressed the same after "
                "it was read over and explained.\n\n"
                "Signature of Accused: ____________________\n"
                "Signature of Magistrate: ____________________"
            ),
            "version": 1,
        },
    ]

    for d in defs:
        d["placeholders"] = extract_placeholders(d["template_body"])
        tpl = DocumentTemplate(
            id=d["id"],
            template_name=d["template_name"],
            category=d["category"],
            template_body=d["template_body"],
            placeholders=d["placeholders"],
            is_active=True,
            version=d["version"],
        )
        db.add(tpl)
    await db.flush()


# ---------------------------------------------------------------------------
# Template queries
# ---------------------------------------------------------------------------

async def list_templates(
    category: Optional[str] = None,
    *,
    db: AsyncSession,
) -> List[dict]:
    """Return templates, optionally filtered by *category*."""
    stmt = select(DocumentTemplate).where(DocumentTemplate.is_active == True)  # noqa: E712
    if category:
        stmt = stmt.where(DocumentTemplate.category == category)
    result = await db.execute(stmt)
    return [_template_to_dict(t) for t in result.scalars().all()]


async def get_template(template_id: str, *, db: AsyncSession) -> Optional[dict]:
    """Return a single template by ID, or ``None``."""
    tpl = await db.get(DocumentTemplate, template_id)
    return _template_to_dict(tpl) if tpl else None


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

async def generate_document(
    template_id: str,
    case_data: dict,
    user_id: str,
    *,
    db: AsyncSession,
) -> dict:
    """Generate a document by substituting placeholders with *case_data*.

    Missing fields are left as ``{{placeholder}}`` and recorded in
    ``missing_fields`` so the IO can fill them in manually.
    """
    tpl = await db.get(DocumentTemplate, template_id)
    if tpl is None:
        raise ValueError(f"Template '{template_id}' not found.")

    body = tpl.template_body
    placeholders = tpl.placeholders or []

    # Build a safe data dict: missing keys map back to their token form.
    render_data: Dict[str, str] = {}
    auto_filled: List[str] = []
    missing: List[str] = []

    for ph in placeholders:
        if ph in case_data and case_data[ph] not in (None, ""):
            render_data[ph] = str(case_data[ph])
            auto_filled.append(ph)
        else:
            render_data[ph] = "{{" + ph + "}}"
            missing.append(ph)

    # Jinja2 uses {{ }} natively; convert our {{token}} -> {{ token }}
    # so the Template engine can substitute them properly.
    jinja_body = _PLACEHOLDER_RE.sub(r"{{ \1 }}", body)
    rendered = Template(jinja_body).render(**render_data)

    doc = GeneratedDocument(
        template_id=template_id,
        case_id=case_data.get("case_id", ""),
        document_category=tpl.category,
        generated_content=rendered,
        auto_filled_fields={"filled": auto_filled, "missing": missing},
        created_by=user_id,
    )
    _refresh_generated_document_hash(doc)
    db.add(doc)
    await db.flush()
    return _gendoc_to_dict(doc, auto_filled, missing)


# ---------------------------------------------------------------------------
# IO editing
# ---------------------------------------------------------------------------

async def update_generated_document(
    doc_id: str,
    content: str,
    user_id: str,
    *,
    db: AsyncSession,
) -> dict:
    """Update the generated content after IO review / manual edits."""
    doc = await db.get(GeneratedDocument, doc_id)
    if doc is None:
        raise ValueError(f"Generated document '{doc_id}' not found.")
    signature_status = doc.digital_signature_status.value if hasattr(doc.digital_signature_status, "value") else doc.digital_signature_status
    if signature_status == "Signed":
        raise ValueError("Signed documents are read-only and cannot be edited.")

    doc.generated_content = content
    _refresh_generated_document_hash(doc)
    doc.updated_by = user_id
    doc.io_edited = True
    await db.flush()
    af = doc.auto_filled_fields or {}
    return _gendoc_to_dict(doc, af.get("filled", []), af.get("missing", []))


async def sign_generated_document(
    doc_id: str,
    user_id: str,
    pin: Optional[str],
    *,
    db: AsyncSession,
) -> dict:
    """Apply a DSC signature when a token has been detected by configuration."""
    import os

    doc = await db.get(GeneratedDocument, doc_id)
    if doc is None:
        raise ValueError(f"Generated document '{doc_id}' not found.")
    if os.getenv("DSC_TOKEN_PRESENT", "").lower() not in {"1", "true", "yes"}:
        raise RuntimeError("Digital Signature Certificate not detected. Please insert your DSC token and try again.")
    if not pin:
        raise ValueError("DSC PIN is required.")

    doc.digital_signature_status = "Signed"
    _refresh_generated_document_hash(doc)
    doc.signed_by = user_id
    doc.signed_at = datetime.now(timezone.utc)
    doc.signature_certificate_details = {
        "certificate_subject": os.getenv("DSC_CERT_SUBJECT", "Configured DSC token"),
        "certificate_serial": os.getenv("DSC_CERT_SERIAL", "unavailable"),
        "provider": os.getenv("DSC_PROVIDER", "Local DSC bridge"),
        "document_sha256": doc.sha256_hash,
    }
    await db.flush()
    af = doc.auto_filled_fields or {}
    return _gendoc_to_dict(doc, af.get("filled", []), af.get("missing", []))


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

async def export_docx(doc_id: str, *, db: AsyncSession) -> bytes:
    """Create a professionally formatted DOCX and return its bytes."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = await db.get(GeneratedDocument, doc_id)
    if doc is None:
        raise ValueError(f"Generated document '{doc_id}' not found.")

    tpl = await db.get(DocumentTemplate, doc.template_id) if doc.template_id else None
    title = tpl.template_name if tpl else "Generated Document"

    category_val = doc.document_category.value if hasattr(doc.document_category, "value") else doc.document_category
    created_at_str = doc.created_at.isoformat()[:10] if doc.created_at else ""

    docx = DocxDocument()

    # -- page margins
    for section in docx.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # -- title
    heading = docx.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- metadata line
    meta = docx.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Category: {category_val}  |  Generated: {created_at_str}"
    )
    run.font.size = Pt(9)
    run.font.italic = True

    docx.add_paragraph("")  # spacer

    # -- body: split on blank lines to form paragraphs; detect headings
    content = doc.generated_content or ""
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        # Treat single short uppercase lines as sub-headings
        if (
            len(lines) == 1
            and len(lines[0]) < 80
            and lines[0] == lines[0].upper()
        ):
            h = docx.add_heading(lines[0], level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para = docx.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            for i, line in enumerate(lines):
                if i > 0:
                    para.add_run("\n")
                # Bold lines that look like field labels (contain a colon)
                if ":" in line and line.index(":") < 40:
                    label, _, rest = line.partition(":")
                    bold_run = para.add_run(label + ":")
                    bold_run.bold = True
                    bold_run.font.size = Pt(11)
                    normal_run = para.add_run(rest)
                    normal_run.font.size = Pt(11)
                else:
                    r = para.add_run(line)
                    r.font.size = Pt(11)

    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

async def export_pdf(doc_id: str, *, db: AsyncSession) -> bytes:
    """Create a PDF rendition of the generated document and return bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = await db.get(GeneratedDocument, doc_id)
    if doc is None:
        raise ValueError(f"Generated document '{doc_id}' not found.")

    tpl = await db.get(DocumentTemplate, doc.template_id) if doc.template_id else None
    title = tpl.template_name if tpl else "Generated Document"

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontSize=16,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "DocBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=15,
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "DocHeading",
        parent=styles["Heading2"],
        fontSize=13,
        spaceAfter=6,
        spaceBefore=12,
    )

    story: List = []
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.3 * cm))

    # Escape XML entities for ReportLab paragraphs.
    def _esc(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    content = doc.generated_content or ""
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        if (
            len(lines) == 1
            and len(lines[0]) < 80
            and lines[0] == lines[0].upper()
        ):
            story.append(Paragraph(_esc(lines[0]), heading_style))
        else:
            text = "<br/>".join(_esc(line) for line in lines)
            story.append(Paragraph(text, body_style))

    pdf.build(story)
    return buf.getvalue()
