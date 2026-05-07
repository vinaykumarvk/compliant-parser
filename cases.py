from __future__ import annotations

"""Case management service layer for the IQW platform.

Async SQLAlchemy ORM implementation.  Every public function returns plain
dicts so the API layer can serialise them directly.
"""

import asyncio
import json
import re
import uuid as _uuid
from datetime import date, datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

import sqlalchemy as sa
from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from audit import compute_sha256, log_audit_event
from cctns import CCTNSSyncResult, sync_case_metadata
from external_interfaces import (
    ExternalServiceError,
    ExternalServiceUnavailable,
    LiveLLMClient,
    delete_object,
    get_object_bytes,
    infer_mime_type,
    is_document_ai_supported_mime,
    put_object_bytes,
    run_configured_ocr,
)
from ocr_enhancements import classify_confidence, detect_language_enhanced
from database import parse_records
from models import (
    Case, CaseDocument, CaseActivity, ActionTrackerTask,
    Notification, PoliceStation, OffenceType, DocumentType, User,
)


# ---------------------------------------------------------------------------
# Lazy import helper (avoids circular deps with api_v1)
# ---------------------------------------------------------------------------

def _api_error(code: str, message: str, field: Optional[str] = None) -> None:
    status = {
        "VALIDATION_ERROR": 400,
        "AUTHENTICATION_ERROR": 401,
        "AUTHORIZATION_ERROR": 403,
        "ACCOUNT_LOCKED": 423,
        "NOT_FOUND": 404,
        "CONFLICT": 409,
        "RATE_LIMITED": 429,
        "SERVER_ERROR": 500,
        "SERVICE_UNAVAILABLE": 503,
    }.get(code, 500)
    raise HTTPException(
        status_code=status,
        detail={
            "error": {
                "code": code,
                "message": message,
                "field": field,
                "request_id": str(_uuid.uuid4())[:8],
            }
        },
    )


# ---------------------------------------------------------------------------
# Deprecated in-memory stubs — kept for backward compat with tests that
# import them.  No longer used by service functions.
# ---------------------------------------------------------------------------

_cases: Dict[str, dict] = {}               # DEPRECATED — use ORM
_case_documents: Dict[str, dict] = {}      # DEPRECATED — use ORM
_case_activities: Dict[str, dict] = {}     # DEPRECATED — use ORM
_action_tracker_tasks: Dict[str, dict] = {}  # DEPRECATED — use ORM
_notifications: Dict[str, dict] = {}       # DEPRECATED — use ORM
_police_stations: Dict[str, dict] = {}     # DEPRECATED — use ORM
_offence_types: Dict[str, dict] = {}       # DEPRECATED — use ORM


# ---------------------------------------------------------------------------
# Status state machine
# ---------------------------------------------------------------------------

VALID_TRANSITIONS: Dict[str, List[str]] = {
    "Complaint_Received": ["FIR_Registered", "Closed_No_FIR"],
    "FIR_Registered": ["Under_Investigation"],
    "Under_Investigation": ["Charge_Sheet_Filed", "Closure_Report_Filed", "Transferred"],
    "Charge_Sheet_Filed": ["Court_Proceedings"],
    "Closure_Report_Filed": ["Disposed"],
    "Court_Proceedings": ["Disposed"],
}
# Terminal states: Disposed, Transferred, Closed_No_FIR

CASE_STAGE_GUIDANCE: Dict[str, dict] = {
    "Complaint_Received": {
        "label": "Complaint Received",
        "description": "Complaint/petition has been received and recorded.",
        "expected_documents": ["Petition", "ID/address proof if required", "Initial supporting material"],
        "expected_actions": [
            "Register petition with station-scoped number and date",
            "Use AI intake to summarize facts and suggest offence type",
            "Verify complainant, occurrence place, date/time, property loss, injuries, and accused particulars",
            "Assess cognizability and jurisdiction",
            "Make GD/station diary entry",
        ],
        "next_hint": "Register FIR if a cognizable offence is disclosed. If not, record reasons and close or transfer as per jurisdiction.",
    },
    "FIR_Registered": {
        "label": "FIR Registered",
        "description": "FIR has been registered with crime number and offence sections.",
        "expected_documents": ["FIR", "Petition", "Station diary extract"],
        "expected_actions": [
            "Assign IO and supervisory officer",
            "Confirm BNS/BNSS and special-law sections with cited legal basis",
            "Sync FIR metadata with CCTNS",
            "Issue complainant copy and acknowledgement",
            "Create investigation plan and statutory deadline tracker",
        ],
        "next_hint": "Move to investigation after IO assignment, section verification, CCTNS sync, and initial plan creation.",
    },
    "Under_Investigation": {
        "label": "Under Investigation",
        "description": "Active investigation by the Investigating Officer.",
        "expected_documents": ["Witness_Statement", "Seizure_Memo", "Arrest_Memo", "Medical_Report", "FSL_Report", "CDR"],
        "expected_actions": [
            "Visit scene of crime and prepare scene observation/panchanama",
            "Record victim, complainant, and witness statements under applicable BNSS provisions",
            "Collect, seal, seize, photograph, hash, and chain-of-custody tag physical/digital evidence",
            "Use AI congruence checks to compare petition, FIR, statements, and medical/FSL documents",
            "Arrest or issue notice only after legal necessity review and supervisory approval where required",
            "Send exhibits to FSL and track acknowledgements",
            "Obtain CDR, CCTV, bank, platform, or device evidence with proper authorization",
            "Submit progress notes and supervisor review at configured intervals",
        ],
        "next_hint": "File charge sheet within 60/90 days, or submit closure report if evidence insufficient.",
    },
    "Charge_Sheet_Filed": {
        "label": "Charge Sheet Filed",
        "description": "Charge sheet prepared and filed before the court.",
        "expected_documents": ["Charge_Sheet", "Witness_Statement", "FSL_Report"],
        "expected_actions": [
            "Verify legal ingredients against collected evidence",
            "Prepare accused, witness, document, property, and expert lists",
            "Generate evidence certificates for electronic records",
            "Run final AI quality and contradiction checks before filing",
            "Record prosecutor/supervisor review",
        ],
        "next_hint": "Case moves to court proceedings after filing and court cognizance.",
    },
    "Closure_Report_Filed": {
        "label": "Closure Report Filed",
        "description": "Investigation concluded \u2014 closure report submitted to court.",
        "expected_documents": ["Other"],
        "expected_actions": [
            "Prepare closure report category and legal reasons",
            "Notify complainant where required",
            "Attach final investigation summary and evidence index",
            "Submit to Magistrate and track acceptance/protest petition",
        ],
        "next_hint": "Awaiting court decision on closure.",
    },
    "Court_Proceedings": {
        "label": "Court Proceedings",
        "description": "Case is before the court \u2014 trial in progress.",
        "expected_documents": [],
        "expected_actions": [
            "Track summons, witness attendance, exhibits, and court directions",
            "Coordinate with prosecutor for evidence presentation",
            "Record adjournment reasons and next hearing tasks",
            "Update disposal outcome after judgment/order",
        ],
        "next_hint": "Case will be disposed upon judgment.",
    },
    "Transferred": {
        "label": "Transferred",
        "description": "Case transferred to another police station or agency.",
        "expected_documents": [],
        "expected_actions": ["Record transfer order, receiving agency/station, handover date, and document inventory"],
        "next_hint": "Terminal state for this jurisdiction after acknowledged handover.",
    },
    "Disposed": {
        "label": "Disposed",
        "description": "Case concluded \u2014 judgment delivered or closure accepted.",
        "expected_documents": [],
        "expected_actions": ["Archive final order, property disposal, appeal status, and record-retention category"],
        "next_hint": "Terminal state. Use read-only access except authorized corrections.",
    },
    "Closed_No_FIR": {
        "label": "Closed (No FIR)",
        "description": "Complaint assessed \u2014 no cognizable offence disclosed. No FIR registered.",
        "expected_documents": [],
        "expected_actions": ["Record reason for non-registration", "Notify complainant", "Preserve petition and decision note"],
        "next_hint": "Terminal state unless reopened by competent authority or fresh material.",
    },
}


# ---------------------------------------------------------------------------
# Validation helpers
# ---------------------------------------------------------------------------

_CRIME_NO_RE = re.compile(r"^\d{4}/\d{4}$")
_PETITION_NO_RE = re.compile(r"^PET/[A-Z]{2,10}/\d{4}/\d{4}$")
_PETITION_NO_PARTS_RE = re.compile(r"^PET/(?P<station>[A-Z0-9]{2,10})/(?P<year>\d{4})/(?P<seq>\d{4})$")
MAX_CASE_UPLOAD_FILES = 20
MAX_DOCUMENT_SIZE_BYTES = 50 * 1024 * 1024
SUPPORTED_UPLOAD_EXTENSIONS = {".pdf", ".jpeg", ".jpg", ".png", ".docx", ".txt"}
SUPPORTED_UPLOAD_ERROR = "File type not supported. Accepted formats: PDF, JPEG, PNG, DOCX, TXT."
DOCUMENT_SIZE_ERROR = "File exceeds 50 MB limit."
GENERIC_CHECKLIST_NOTE = "Generic checklist applied. Contact AI Admin for offence-specific checklists."
DOCUMENT_TYPE_ALIASES = {
    "Complaint": "Petition",
    "Statement": "Witness_Statement",
    "Evidence": "Other",
    "Report": "Other",
}


def validate_crime_no(crime_no: str) -> bool:
    """Return True when *crime_no* matches NNNN/YYYY."""
    return bool(_CRIME_NO_RE.match(crime_no))


def validate_petition_no(petition_no: str) -> bool:
    """Return True when *petition_no* matches PET/PS/YYYY/NNNN."""
    return bool(_PETITION_NO_RE.match(petition_no))


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _crime_year(crime_no: str) -> Optional[str]:
    if not crime_no or "/" not in crime_no:
        return None
    return crime_no.rsplit("/", 1)[-1]


def _coerce_registration_datetime(value: Any = None) -> datetime:
    """Return a timezone-aware registration datetime for case numbering."""
    if isinstance(value, datetime):
        if value.tzinfo is None:
            return value.replace(tzinfo=timezone.utc)
        return value
    if isinstance(value, date):
        return datetime(value.year, value.month, value.day, tzinfo=timezone.utc)
    if isinstance(value, str) and value.strip():
        cleaned = value.strip().replace("Z", "+00:00")
        try:
            if "T" in cleaned:
                parsed = datetime.fromisoformat(cleaned)
            else:
                parsed_date = date.fromisoformat(cleaned)
                parsed = datetime(parsed_date.year, parsed_date.month, parsed_date.day, tzinfo=timezone.utc)
            return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
        except ValueError:
            _api_error("VALIDATION_ERROR", "Invalid registration date.", "date_of_registration")
    return datetime.now(timezone.utc)


def _safe_station_code(station: Optional[PoliceStation]) -> str:
    code = (getattr(station, "station_code", None) or getattr(station, "id", "") or "PS").upper()
    code = re.sub(r"[^A-Z0-9]", "", code)
    return (code or "PS")[:10]


async def _get_required_station(police_station_id: Optional[str], db: AsyncSession) -> PoliceStation:
    if not police_station_id:
        _api_error("VALIDATION_ERROR", "Police Station is required for case numbering.", "police_station_id")
    station = await db.get(PoliceStation, police_station_id)
    if station is None:
        _api_error("VALIDATION_ERROR", "Selected Police Station was not found.", "police_station_id")
    return station


async def _next_station_sequence(
    *,
    police_station_id: str,
    case_type: str,
    year: int,
    station_code: str,
    db: AsyncSession,
) -> int:
    result = await db.execute(
        select(Case).where(
            Case.police_station_id == police_station_id,
            Case.is_deleted == False,  # noqa: E712
        )
    )
    max_seq = 0
    for row in result.scalars().all():
        if case_type == "FIR":
            crime_no = row.crime_no or ""
            if _crime_year(crime_no) != str(year):
                continue
            seq_text = crime_no.split("/", 1)[0]
        else:
            match = _PETITION_NO_PARTS_RE.match(row.petition_no or "")
            if not match or match.group("station") != station_code or match.group("year") != str(year):
                continue
            seq_text = match.group("seq")
        try:
            max_seq = max(max_seq, int(seq_text))
        except ValueError:
            continue
    return max_seq + 1


async def generate_case_identifiers(
    *,
    police_station_id: str,
    case_type: str,
    db: AsyncSession,
    registration_datetime: Any = None,
) -> dict[str, Any]:
    """Return station-scoped case numbers and registration date for intake."""
    normalized_case_type = case_type or "Petition"
    if normalized_case_type not in {"Petition", "FIR", "Suo_Motu"}:
        _api_error("VALIDATION_ERROR", "Invalid case type.", "case_type")
    station = await _get_required_station(police_station_id, db)
    registered_at = _coerce_registration_datetime(registration_datetime)
    station_code = _safe_station_code(station)
    year = registered_at.year
    payload: dict[str, Any] = {
        "case_type": normalized_case_type,
        "police_station_id": police_station_id,
        "police_station_name": station.name,
        "station_code": station_code,
        "date_of_registration": registered_at.isoformat(),
        "crime_no": None,
        "petition_no": None,
    }
    if normalized_case_type in {"FIR", "Suo_Motu"}:
        seq = await _next_station_sequence(
            police_station_id=police_station_id,
            case_type="FIR",
            year=year,
            station_code=station_code,
            db=db,
        )
        payload["crime_no"] = f"{seq:04d}/{year}"
    if normalized_case_type == "Petition":
        seq = await _next_station_sequence(
            police_station_id=police_station_id,
            case_type="Petition",
            year=year,
            station_code=station_code,
            db=db,
        )
        payload["petition_no"] = f"PET/{station_code}/{year}/{seq:04d}"
    return payload


def _enum_values(enum_cls) -> set[str]:
    return {item.value for item in enum_cls}


def _normalise_task_status(value: str) -> str:
    lookup = {
        "done": "Completed",
        "completed": "Completed",
        "pending": "Pending",
        "overdue": "Overdue",
    }
    return lookup.get(str(value).strip().lower(), value)


def _normalise_priority(value: str) -> str:
    lookup = {"high": "High", "medium": "Medium", "low": "Low"}
    return lookup.get(str(value).strip().lower(), value)


GLOBAL_CASE_ACCESS_ROLES = {"System_Admin", "AI_Admin", "Admin", "Supervisor"}


def can_access_case(case: Case, user: dict) -> bool:
    """Return True when the authenticated user may access the case."""
    role = user.get("role")
    user_id = user.get("sub")
    if role in GLOBAL_CASE_ACCESS_ROLES:
        return True
    return bool(user_id and user_id in {case.io_id, case.created_by})


async def ensure_case_access(case_id: str, user: dict, db: AsyncSession) -> Case:
    """Load a case and enforce case-level authorization for direct-ID routes."""
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")
    if not can_access_case(case, user):
        _api_error("AUTHORIZATION_ERROR", f"User is not authorized to access case '{case_id}'.")
    return case


# ---------------------------------------------------------------------------
# ORM → dict helper
# ---------------------------------------------------------------------------

def _to_dict(obj, key_map: Optional[dict] = None, exclude: Optional[set] = None) -> dict:
    """Convert ORM model to plain dict, applying optional key renames."""
    from sqlalchemy import inspect as sa_inspect
    result = {}
    for attr in sa_inspect(type(obj)).column_attrs:
        key = attr.key
        if exclude and key in exclude:
            continue
        val = getattr(obj, key)
        if isinstance(val, datetime):
            val = val.isoformat()
        elif hasattr(val, 'value'):  # Enum
            val = val.value
        elif isinstance(val, date) and not isinstance(val, datetime):
            val = val.isoformat()
        out_key = key_map.get(key, key) if key_map else key
        result[out_key] = val
    return result


def _case_to_dict(case: Case, sync_result: Optional[CCTNSSyncResult] = None) -> dict:
    data = _to_dict(case)
    # Include human-readable names from relationships (if loaded)
    try:
        ps = case.police_station
        data["police_station_name"] = ps.name if ps else None
    except Exception:
        data["police_station_name"] = None
    try:
        ot = case.primary_offence_type
        data["primary_offence_type_name"] = ot.name if ot else None
    except Exception:
        data["primary_offence_type_name"] = None
    status = data.get("cctns_sync_status")
    data["cctns_retry_available"] = status in ("Pending", "Failed")
    if sync_result:
        data["cctns_sync_attempts"] = sync_result.attempts
        data["cctns_sync_message"] = sync_result.error
        data["cctns_retry_available"] = sync_result.retry_available
    return data


def _task_to_dict(task: ActionTrackerTask, *, today: Optional[date] = None) -> dict:
    data = _to_dict(task)
    today = today or datetime.now(timezone.utc).date()
    due_raw = task.due_date
    due_date = due_raw if isinstance(due_raw, date) else None
    status = data.get("status")
    snoozed_until = task.snoozed_until
    snoozed_active = bool(
        snoozed_until and isinstance(snoozed_until, datetime) and snoozed_until > datetime.now(timezone.utc)
    )
    is_completed = status == "Completed"
    is_overdue = bool(due_date and due_date < today and not is_completed and not snoozed_active)
    if is_overdue:
        data["status"] = "Overdue"
    data["is_overdue"] = is_overdue
    data["reminder_due"] = bool(due_date and (due_date - today).days in (3, 1, 0) and not is_completed)
    return data


async def _ensure_unique_crime_no(crime_no: str, police_station_id: Optional[str], db: AsyncSession) -> None:
    if not crime_no:
        return
    year = _crime_year(crime_no)
    stmt = select(Case).where(
        Case.crime_no == crime_no,
        Case.police_station_id == police_station_id,
        Case.is_deleted == False,  # noqa: E712
    )
    result = await db.execute(stmt)
    duplicate = result.scalars().first()
    if duplicate and _crime_year(duplicate.crime_no or "") == year:
        _api_error(
            "CONFLICT",
            f"Crime Number {crime_no} already exists for this police station.",
            "crime_no",
        )


async def _match_offence_type(name: Optional[str], db: AsyncSession) -> Optional[OffenceType]:
    if not name:
        return None
    normalized = name.strip().lower()
    if not normalized:
        return None
    result = await db.execute(
        select(OffenceType).where(OffenceType.is_active == True)  # noqa: E712
    )
    offences = result.scalars().all()
    for offence in offences:
        if (offence.name or "").strip().lower() == normalized:
            return offence
    for offence in offences:
        if normalized in (offence.name or "").strip().lower() or (offence.name or "").strip().lower() in normalized:
            return offence
    return None


async def _offence_catalog_for_prompt(db: AsyncSession) -> list[dict[str, Any]]:
    result = await db.execute(
        select(OffenceType).where(OffenceType.is_active == True)  # noqa: E712
    )
    catalog = []
    for row in result.scalars().all():
        catalog.append(
            {
                "id": row.id,
                "name": row.name,
                "category": row.category,
                "bns_section": row.bns_section,
                "ipc_section": row.ipc_section,
            }
        )
    return catalog


def _extract_text_for_intake(file_name: str, file_bytes: bytes, mime_type: Optional[str]) -> dict[str, Any]:
    detected_mime = infer_mime_type(file_name, mime_type)
    lower_name = (file_name or "").lower()
    if lower_name.endswith(".txt") or detected_mime.startswith("text/"):
        return {
            "text": file_bytes.decode("utf-8", errors="replace"),
            "ocr": {
                "provider": "direct_text",
                "mime_type": detected_mime,
                "page_count": None,
            },
        }
    if lower_name.endswith(".docx") or detected_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document

            document = Document(BytesIO(file_bytes))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        except Exception as exc:
            raise ExternalServiceError(f"DOCX petition extraction failed: {exc}") from exc
        return {
            "text": text,
            "ocr": {
                "provider": "docx_text",
                "mime_type": detected_mime,
                "page_count": None,
            },
        }
    if not is_document_ai_supported_mime(detected_mime):
        _api_error(
            "VALIDATION_ERROR",
            "Petition AI intake supports PDF, image, DOCX, or TXT uploads.",
            "petition_file",
        )
    ocr_result = run_configured_ocr(
        file_bytes,
        file_name=file_name,
        mime_type=detected_mime,
    )
    return {
        "text": ocr_result.text or "",
        "ocr": {
            "provider": ocr_result.provider,
            "mime_type": detected_mime,
            "page_count": ocr_result.page_count,
            "confidence": ocr_result.confidence,
        },
    }


async def suggest_case_intake_from_petition(
    *,
    file_name: str,
    file_bytes: bytes,
    mime_type: Optional[str],
    user_id: str,
    db: AsyncSession,
) -> dict[str, Any]:
    """Use OCR and a live LLM to produce editable case-intake suggestions."""
    if not file_bytes:
        _api_error("VALIDATION_ERROR", "Petition file is empty.", "petition_file")

    try:
        extracted = await asyncio.to_thread(_extract_text_for_intake, file_name, file_bytes, mime_type)
    except ExternalServiceUnavailable:
        raise
    except ExternalServiceError:
        raise

    text = " ".join((extracted.get("text") or "").split())
    if not text:
        _api_error("VALIDATION_ERROR", "No readable text was found in the petition.", "petition_file")

    offence_catalog = await _offence_catalog_for_prompt(db)
    system_prompt = (
        "You are a senior Indian police case-intake officer and legal AI assistant. "
        "Read the petition text and return JSON only. Do not invent facts. "
        "Choose the closest offence_type from the supplied offence catalog when possible. "
        "Write a neutral brief_facts summary suitable for a case register. "
        "If the petition does not disclose a cognizable offence, say so in risk_flags."
    )
    user_prompt = json.dumps(
        {
            "petition_text": text,
            "offence_catalog": offence_catalog,
            "required_json_schema": {
                "brief_facts": "2-4 concise sentences, no legal conclusions unless explicit",
                "offence_type": "one offence catalog name or null",
                "offence_confidence": "High, Medium, or Low",
                "case_type": "Petition, FIR, or Suo_Motu",
                "date_of_occurrence": "YYYY-MM-DD or null",
                "risk_flags": ["strings"],
                "rationale": "short non-PII reason for the suggestion",
            },
        },
        ensure_ascii=False,
    )
    llm_result = await asyncio.to_thread(
        LiveLLMClient().generate_json,
        system_prompt,
        user_prompt,
        timeout=180,
    )
    data = llm_result.data if isinstance(llm_result.data, dict) else {}
    offence = await _match_offence_type(data.get("offence_type"), db)
    offence_name = offence.name if offence else (data.get("offence_type") or None)
    suggestion_id = str(_uuid.uuid4())
    payload = {
        "suggestion_id": suggestion_id,
        "brief_facts": data.get("brief_facts") or "",
        "offence_type": offence_name,
        "primary_offence_type_id": offence.id if offence else None,
        "offence_confidence": data.get("offence_confidence") or "Low",
        "case_type": data.get("case_type") if data.get("case_type") in {"Petition", "FIR", "Suo_Motu"} else "Petition",
        "date_of_occurrence": data.get("date_of_occurrence"),
        "risk_flags": data.get("risk_flags") if isinstance(data.get("risk_flags"), list) else [],
        "rationale": data.get("rationale") or "",
        "editable_fields": ["brief_facts", "primary_offence_type_id", "offence_type"],
        "ocr": extracted["ocr"] | {"text_length": len(text)},
        "llm": {
            "provider": llm_result.provider,
            "model": llm_result.model,
            "privacy": llm_result.privacy,
        },
    }
    await log_audit_event(
        user_id=user_id,
        action_type="AI_Analysis",
        entity_type="CaseIntakeSuggestion",
        entity_id=suggestion_id,
        details={
            "source": "petition_intake",
            "file_name": Path(file_name or "petition").name,
            "sha256": compute_sha256(file_bytes),
            "ocr_provider": payload["ocr"]["provider"],
            "llm_provider": llm_result.provider,
            "llm_model": llm_result.model,
            "privacy": llm_result.privacy,
            "suggested_fields": ["brief_facts", "offence_type", "primary_offence_type_id"],
        },
        db=db,
    )
    return payload


def _value_audit_summary(value: Any) -> dict[str, Any]:
    if value is None:
        return {"present": False}
    if isinstance(value, str):
        cleaned = value.strip()
        return {
            "present": bool(cleaned),
            "length": len(cleaned),
            "sha256": compute_sha256(cleaned.encode("utf-8")) if cleaned else None,
        }
    encoded = json.dumps(value, sort_keys=True, default=str).encode("utf-8")
    return {
        "present": True,
        "type": type(value).__name__,
        "sha256": compute_sha256(encoded),
    }


def _ai_suggestion_diffs(data: dict) -> list[dict[str, Any]]:
    context = data.get("ai_suggestion_context")
    if not isinstance(context, dict):
        return []
    mapping = {
        "brief_facts": "suggested_brief_facts",
        "primary_offence_type_id": "suggested_primary_offence_type_id",
        "offence_type": "suggested_offence_type",
    }
    diffs = []
    for field, suggested_key in mapping.items():
        if field not in data:
            continue
        suggested = context.get(suggested_key)
        submitted = data.get(field)
        if (suggested or None) != (submitted or None):
            diffs.append(
                {
                    "field": field,
                    "suggested": _value_audit_summary(suggested),
                    "submitted": _value_audit_summary(submitted),
                }
            )
    return diffs


def validate_document_upload(
    file_name: str,
    document_type: str,
    file_bytes: bytes,
) -> None:
    """Validate BRD upload constraints before a CaseDocument is persisted."""
    if not document_type:
        _api_error("VALIDATION_ERROR", "Document type is required before upload completes.", "document_type")
    normalized_type = DOCUMENT_TYPE_ALIASES.get(document_type, document_type)
    if normalized_type not in _enum_values(DocumentType):
        _api_error("VALIDATION_ERROR", f"Invalid document type '{document_type}'.", "document_type")
    if len(file_bytes) > MAX_DOCUMENT_SIZE_BYTES:
        _api_error("VALIDATION_ERROR", DOCUMENT_SIZE_ERROR, "file")
    lower_name = (file_name or "").lower()
    if not any(lower_name.endswith(ext) for ext in SUPPORTED_UPLOAD_EXTENSIONS):
        _api_error("VALIDATION_ERROR", SUPPORTED_UPLOAD_ERROR, "file")


def normalize_document_type(document_type: str) -> str:
    return DOCUMENT_TYPE_ALIASES.get(document_type, document_type)


async def read_case_document_bytes(doc: CaseDocument) -> Optional[bytes]:
    if getattr(doc, "file_storage_uri", None):
        return await get_object_bytes(doc.file_storage_uri)
    if doc.file_bytes is not None:
        return doc.file_bytes
    if doc.file_path:
        try:
            return Path(doc.file_path).read_bytes()
        except OSError:
            return None
    return None


# Key-rename maps for models whose ORM column names differ from the legacy
# dict keys that the API layer expects.

_DOCUMENT_KEY_MAP = {
    "sha256_hash": "sha256",
    "file_size_bytes": "size_bytes",
    "created_by": "uploaded_by",
    "created_at": "uploaded_at",
}

_DOCUMENT_EXCLUDE_KEYS = {"file_bytes"}

_POLICE_STATION_KEY_MAP = {
    "station_code": "code",
}


# ---------------------------------------------------------------------------
# Case CRUD
# ---------------------------------------------------------------------------

async def create_case(data: dict, user_id: str, db: AsyncSession) -> dict:
    """Create a new case after validating crime/petition numbers."""
    case_type = data.get("case_type", "FIR")
    police_station_id = data.get("police_station_id")
    date_of_registration = _coerce_registration_datetime(data.get("date_of_registration"))
    crime_no = data.get("crime_no", "")
    petition_no = data.get("petition_no", "")

    if case_type in {"FIR", "Suo_Motu"} and not crime_no:
        generated = await generate_case_identifiers(
            police_station_id=police_station_id,
            case_type=case_type,
            db=db,
            registration_datetime=date_of_registration,
        )
        crime_no = generated.get("crime_no") or ""
    if case_type == "Petition" and not petition_no:
        generated = await generate_case_identifiers(
            police_station_id=police_station_id,
            case_type=case_type,
            db=db,
            registration_datetime=date_of_registration,
        )
        petition_no = generated.get("petition_no") or ""

    if crime_no and not validate_crime_no(crime_no):
        _api_error("VALIDATION_ERROR", "Invalid Crime Number format. Expected: NNNN/YYYY.", "crime_no")

    if petition_no and not validate_petition_no(petition_no):
        _api_error("VALIDATION_ERROR", "Invalid Petition Number format. Expected: PET/PS/YYYY/NNNN.", "petition_no")

    if case_type == "FIR" and not crime_no:
        _api_error("VALIDATION_ERROR", "Crime Number is required for FIR cases.", "crime_no")
    if case_type == "Petition" and not petition_no:
        _api_error("VALIDATION_ERROR", "Petition Number is required for petition cases.", "petition_no")

    await _ensure_unique_crime_no(crime_no, police_station_id, db)
    ai_diffs = _ai_suggestion_diffs(
        {
            **data,
            "crime_no": crime_no,
            "petition_no": petition_no,
        }
    )

    # Extract petition analysis from AI suggestion context for persistence
    ai_ctx = data.get("ai_suggestion_context")
    petition_analysis_payload = None
    if isinstance(ai_ctx, dict):
        petition_analysis_payload = {
            "suggestion_id": ai_ctx.get("suggestion_id"),
            "brief_facts": ai_ctx.get("brief_facts"),
            "offence_type": ai_ctx.get("offence_type"),
            "offence_confidence": ai_ctx.get("offence_confidence"),
            "case_type": ai_ctx.get("case_type"),
            "date_of_occurrence": ai_ctx.get("date_of_occurrence"),
            "risk_flags": ai_ctx.get("risk_flags", []),
            "rationale": ai_ctx.get("rationale"),
            "extracted_at": datetime.now(timezone.utc).isoformat(),
        }

    case = Case(
        case_type=case_type,
        crime_no=crime_no,
        petition_no=petition_no,
        brief_facts=data.get("brief_facts", ""),
        offence_type=data.get("offence_type", ""),
        police_station_id=police_station_id,
        date_of_registration=date_of_registration,
        primary_offence_type_id=data.get("primary_offence_type_id"),
        secondary_offence_type_ids=data.get("secondary_offence_type_ids", []),
        petition_analysis=petition_analysis_payload,
        status="Complaint_Received",
        cctns_sync_status="Pending",
        created_by=user_id,
        io_id=data.get("io_id", user_id),
    )
    db.add(case)
    await db.flush()  # assigns id

    await add_activity(case.id, "Case_Created", user_id, "Case registered and opened.", db=db)
    context = data.get("ai_suggestion_context")
    if isinstance(context, dict):
        await add_activity(
            case.id,
            "AI_Intake_Applied",
            user_id,
            "AI petition intake suggestions were reviewed during case creation.",
            entity_type="case",
            entity_id=case.id,
            db=db,
        )
        await log_audit_event(
            user_id=user_id,
            action_type="Edit" if ai_diffs else "AI_Analysis",
            entity_type="Case",
            entity_id=case.id,
            details={
                "source": "ai_petition_intake",
                "suggestion_id": context.get("suggestion_id"),
                "provider": context.get("provider"),
                "model": context.get("model"),
                "editable_fields_reviewed": ["brief_facts", "offence_type", "primary_offence_type_id"],
                "overrides": ai_diffs,
            },
            db=db,
        )
    sync_result = await sync_case_metadata(_case_to_dict(case))
    case.cctns_sync_status = sync_result.status
    if sync_result.cctns_case_id:
        case.cctns_case_id = sync_result.cctns_case_id
    await db.flush()

    await add_activity(
        case.id,
        "CCTNS_Sync",
        user_id,
        sync_result.error or f"CCTNS sync status: {sync_result.status}.",
        entity_type="case",
        entity_id=case.id,
        db=db,
    )
    if sync_result.status == "Failed":
        await create_notification(
            user_id,
            "warning",
            sync_result.error or f"CCTNS sync failed for case {crime_no} after 3 attempts. Manual retry available.",
            entity_type="case",
            entity_id=case.id,
            db=db,
        )
    await db.refresh(case, ["police_station", "primary_offence_type"])
    return _case_to_dict(case, sync_result)


async def get_case(case_id: str, db: AsyncSession) -> Optional[dict]:
    """Return a single case or None."""
    stmt = (
        select(Case)
        .where(Case.id == case_id)
        .options(
            selectinload(Case.police_station),
            selectinload(Case.primary_offence_type),
        )
    )
    result = await db.execute(stmt)
    case = result.scalar_one_or_none()
    return _case_to_dict(case) if case else None


async def list_cases(user_id: str, role: str, filters: Optional[dict] = None, *, db: AsyncSession) -> dict:
    """List cases visible to the caller with pagination.

    IO sees only own cases; System_Admin / Admin sees all.
    Filters: status, police_station_id, date_from, date_to, page, page_size.
    Returns: {"items": [...], "total": N, "page": P, "page_size": S}
    """
    stmt = (
        select(Case)
        .where(Case.is_deleted == False)  # noqa: E712
        .options(
            selectinload(Case.police_station),
            selectinload(Case.primary_offence_type),
        )
    )

    if role in ("System_Admin", "Admin", "AI_Admin"):
        pass  # full visibility
    elif role in ("Clerk", "SHO"):
        # Clerk / SHO see all cases at their police station
        user_obj = await db.get(User, user_id)
        if user_obj and user_obj.police_station_id:
            stmt = stmt.where(Case.police_station_id == user_obj.police_station_id)
        else:
            # No station assigned — fall back to own cases only
            stmt = stmt.where(Case.io_id == user_id)
    else:
        # IO and others see only their own cases
        stmt = stmt.where(Case.io_id == user_id)

    if filters:
        if filters.get("status"):
            stmt = stmt.where(Case.status == filters["status"])
        if filters.get("police_station_id"):
            stmt = stmt.where(Case.police_station_id == filters["police_station_id"])
        if filters.get("date_from"):
            stmt = stmt.where(Case.created_at >= filters["date_from"])
        if filters.get("date_to"):
            stmt = stmt.where(Case.created_at <= filters["date_to"])

    stmt = stmt.order_by(Case.created_at.desc())

    # total count
    count_stmt = select(sa.func.count()).select_from(stmt.subquery())
    total = (await db.execute(count_stmt)).scalar() or 0

    # paginate
    page = int((filters or {}).get("page", 1))
    page_size = int((filters or {}).get("page_size", 20))
    stmt = stmt.offset((page - 1) * page_size).limit(page_size)

    result = await db.execute(stmt)
    items = [_case_to_dict(c) for c in result.scalars().all()]

    # Enrich with petitioner name & language from latest document
    case_ids = [item["id"] for item in items]
    if case_ids:
        doc_stmt = (
            select(CaseDocument)
            .where(
                CaseDocument.case_id.in_(case_ids),
                CaseDocument.is_latest_version == True,  # noqa: E712
                CaseDocument.is_deleted == False,  # noqa: E712
                CaseDocument.parsed_output.isnot(None),
            )
            .order_by(CaseDocument.created_at.desc())
        )
        doc_result = await db.execute(doc_stmt)
        doc_map: dict = {}
        for doc in doc_result.scalars().all():
            if doc.case_id not in doc_map:
                doc_map[doc.case_id] = doc
        for item in items:
            doc = doc_map.get(item["id"])
            if doc and doc.parsed_output:
                po = doc.parsed_output if isinstance(doc.parsed_output, dict) else {}
                fir = po.get("fir_draft") or {}
                informant = fir.get("informant") or {}
                parties = fir.get("parties") or {}
                who = parties.get("who_components") or {}
                complainant_vals = (who.get("complainant") or {}).get("values") or []
                raw_name = (
                    informant.get("name")
                    or (complainant_vals[0] if complainant_vals else None)
                    or parties.get("victim")
                    or None
                )
                # Clean up common extraction artefacts
                if raw_name:
                    raw_name = raw_name.rstrip(" (").strip()
                    low = raw_name.lower()
                    # Reject values that are clearly not person names
                    if any(w in low for w in ("registration", "investigation", "complaint", "fir ")):
                        raw_name = None
                    elif len(raw_name) > 60:
                        raw_name = raw_name[:57] + "..."
                item["petitioner_name"] = raw_name or None
                lang = po.get("language") or {}
                item["petition_language"] = lang.get("detected_name") or doc.language_detected or None
            else:
                item["petitioner_name"] = None
                item["petition_language"] = None

    return {"items": items, "total": total, "page": page, "page_size": page_size}


async def update_case(case_id: str, data: dict, user_id: str, db: AsyncSession) -> dict:
    """Update mutable fields on an existing case."""
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    mutable = ("brief_facts", "offence_type", "police_station_id",
               "primary_offence_type_id", "secondary_offence_type_ids")
    before = {key: getattr(case, key) for key in mutable}
    for key in mutable:
        if key in data:
            setattr(case, key, data[key])

    case.updated_at = datetime.now(timezone.utc)
    await db.flush()

    await add_activity(case_id, "Case_Updated", user_id, "Case details updated.", db=db)
    ai_diffs = _ai_suggestion_diffs(data)
    if ai_diffs:
        changed_fields = [
            key
            for key in mutable
            if key in data and (before.get(key) or None) != (data.get(key) or None)
        ]
        await add_activity(
            case_id,
            "AI_Suggestion_Edited",
            user_id,
            "Editable AI-suggested case fields were changed by the user.",
            entity_type="case",
            entity_id=case_id,
            db=db,
        )
        await log_audit_event(
            user_id=user_id,
            action_type="Edit",
            entity_type="Case",
            entity_id=case_id,
            details={
                "source": "ai_petition_intake_edit",
                "suggestion_id": (data.get("ai_suggestion_context") or {}).get("suggestion_id"),
                "changed_fields": changed_fields,
                "overrides": ai_diffs,
            },
            db=db,
        )
    await db.refresh(case, ["police_station", "primary_offence_type"])
    return _case_to_dict(case)


# ---------------------------------------------------------------------------
# Offence type modification (AC-002-4)
# ---------------------------------------------------------------------------


async def update_case_offence_type(
    case_id: str,
    primary_offence_type_id: Optional[str],
    secondary_offence_type_ids: Optional[list] = None,
    *,
    user_id: str,
    db: AsyncSession,
) -> dict:
    """Explicitly update offence classification during investigation.

    Returns the updated case dict.  Raises NOT_FOUND if the case does not exist.
    """
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    before_primary = case.primary_offence_type_id
    before_secondary = case.secondary_offence_type_ids

    case.primary_offence_type_id = primary_offence_type_id
    if secondary_offence_type_ids is not None:
        case.secondary_offence_type_ids = secondary_offence_type_ids
    case.updated_at = datetime.now(timezone.utc)
    await db.flush()

    await add_activity(
        case_id,
        "Offence_Type_Updated",
        user_id,
        f"Primary offence changed from {before_primary} to {primary_offence_type_id}.",
        db=db,
    )
    await db.refresh(case, ["primary_offence_type"])
    return _case_to_dict(case)


# ---------------------------------------------------------------------------
# Status transitions
# ---------------------------------------------------------------------------

async def transition_case_status(case_id: str, new_status: str, user_id: str, db: AsyncSession) -> dict:
    """Move a case to *new_status* if the transition is valid."""
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    current = case.status if isinstance(case.status, str) else case.status.value
    allowed = VALID_TRANSITIONS.get(current, [])
    if new_status not in allowed:
        _api_error(
            "VALIDATION_ERROR",
            f"Cannot transition from '{current}' to '{new_status}'. "
            f"Allowed: {allowed}.",
            "status",
        )

    # Transitioning to FIR_Registered requires crime_no on the case
    if new_status == "FIR_Registered":
        crime_no = case.crime_no if isinstance(case.crime_no, str) else ""
        if not crime_no:
            _api_error(
                "VALIDATION_ERROR",
                "Crime Number must be set on the case before registering FIR.",
                "crime_no",
            )

    case.status = new_status
    case.updated_at = datetime.now(timezone.utc)
    await db.flush()

    await add_activity(
        case_id, "Status_Change", user_id,
        f"Status changed from {current} to {new_status}.",
        db=db,
    )
    await db.refresh(case, ["police_station", "primary_offence_type"])
    return _case_to_dict(case)


def get_stage_guidance(status: str) -> dict:
    """Return guidance data for a given case status."""
    return CASE_STAGE_GUIDANCE.get(status, {})


def get_lifecycle_map() -> dict:
    """Return the full lifecycle map: transitions + guidance for frontend rendering."""
    return {
        "transitions": VALID_TRANSITIONS,
        "guidance": CASE_STAGE_GUIDANCE,
        "design_principles": [
            "Petition intake precedes FIR registration unless FIR data already exists.",
            "AI suggestions assist intake and quality checks but officer-reviewed values remain editable.",
            "Every AI-suggested field edit is audited without storing raw PII in audit details.",
            "CCTNS sync, statutory deadline tracking, document integrity, and supervisor review are lifecycle gates.",
        ],
    }


async def seed_demo_case(user_id: str, db: AsyncSession) -> dict:
    """Create a sample demo case with activities for first-time users."""
    case = Case(
        case_type="Petition",
        crime_no="",
        petition_no="PET/MH/2024/0001",
        brief_facts=(
            "Complainant Shri Ramesh Kumar reports theft of mobile phone "
            "(iPhone 15 Pro, IMEI: 123456789012345) from his person while "
            "travelling in local bus route 42 near Andheri Station on 15-Jan-2024 at approx 18:30 hrs."
        ),
        offence_type="Theft",
        police_station_id=None,
        status="Complaint_Received",
        cctns_sync_status="Pending",
        created_by=user_id,
        io_id=user_id,
    )
    db.add(case)
    await db.flush()

    await add_activity(
        case.id, "Case_Created", user_id,
        "Demo case created \u2014 complaint received for mobile theft.",
        db=db,
    )
    await db.refresh(case, ["police_station", "primary_offence_type"])
    return _case_to_dict(case)


async def retry_cctns_sync(case_id: str, user_id: str, db: AsyncSession) -> dict:
    """Manually retry CCTNS sync for a pending or failed case."""
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    sync_result = await sync_case_metadata(_case_to_dict(case), retry_interval_seconds=0)
    case.cctns_sync_status = sync_result.status
    if sync_result.cctns_case_id:
        case.cctns_case_id = sync_result.cctns_case_id
    case.updated_at = datetime.now(timezone.utc)
    await db.flush()

    await add_activity(
        case_id,
        "CCTNS_Retry",
        user_id,
        sync_result.error or f"CCTNS retry completed with status {sync_result.status}.",
        entity_type="case",
        entity_id=case_id,
        db=db,
    )
    await db.refresh(case, ["police_station", "primary_offence_type"])
    return _case_to_dict(case, sync_result)


# ---------------------------------------------------------------------------
# Document management
# ---------------------------------------------------------------------------

async def attach_document(
    case_id: str,
    file_name: str,
    document_type: str,
    file_bytes: bytes,
    user_id: str,
    db: AsyncSession,
    mime_type: Optional[str] = None,
    upload_method: str = "Drag_Drop",
) -> dict:
    """Attach a document to a case, computing its SHA-256 hash."""
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    validate_document_upload(file_name, document_type, file_bytes)
    document_type = normalize_document_type(document_type)
    lower_name = (file_name or "").lower()
    detected_mime = infer_mime_type(file_name, mime_type)
    extracted_text = ""
    language_detected = None
    ocr_confidence = None
    ocr_status = "Not_Required"
    if lower_name.endswith(".txt"):
        extracted_text = file_bytes.decode("utf-8", errors="replace")
        language_detected = detect_language_enhanced(extracted_text)["language"]
        ocr_confidence = classify_confidence(extracted_text)
        ocr_status = "Completed"
    elif is_document_ai_supported_mime(detected_mime):
        try:
            ocr_result = run_configured_ocr(
                file_bytes,
                file_name=file_name,
                mime_type=detected_mime,
            )
            extracted_text = ocr_result.text
            language_detected = detect_language_enhanced(extracted_text)["language"]
            ocr_confidence = classify_confidence(extracted_text)
            ocr_status = "Completed"
        except ExternalServiceUnavailable:
            ocr_status = "Pending"
        except ExternalServiceError:
            ocr_status = "Failed"
    latest_result = await db.execute(
        select(CaseDocument).where(
            CaseDocument.case_id == case_id,
            CaseDocument.file_name == file_name,
            CaseDocument.is_latest_version == True,  # noqa: E712
        )
    )
    previous_latest = latest_result.scalars().first()
    version = 1
    if previous_latest is not None:
        version = (previous_latest.version or 1) + 1
        previous_latest.is_latest_version = False

    sha256 = compute_sha256(file_bytes)
    safe_name = Path(file_name or "document").name or "document"
    storage = await put_object_bytes(
        f"case-documents/{case_id}/{sha256}-{safe_name}",
        file_bytes,
        detected_mime,
    )
    doc = CaseDocument(
        case_id=case_id,
        file_name=file_name,
        document_type=document_type,
        sha256_hash=sha256,
        file_size_bytes=len(file_bytes),
        mime_type=detected_mime,
        file_bytes=None,
        file_storage_uri=storage.uri,
        file_storage_provider=storage.provider,
        file_encryption_key_ref=storage.encryption_key_ref,
        upload_method=upload_method,
        ocr_status=ocr_status,
        ocr_extracted_text=extracted_text or None,
        ocr_confidence=ocr_confidence,
        language_detected=language_detected,
        version=version,
        is_latest_version=True,
        created_by=user_id,
    )
    db.add(doc)
    await db.flush()

    await add_activity(
        case_id, "Document_Attached", user_id,
        f"Document '{file_name}' attached." if version == 1 else f"Document '{file_name}' uploaded as version {version}.",
        entity_type="document", entity_id=doc.id,
        db=db,
    )
    return _to_dict(doc, key_map=_DOCUMENT_KEY_MAP, exclude=_DOCUMENT_EXCLUDE_KEYS)


async def attach_documents(
    case_id: str,
    files: list[dict],
    user_id: str,
    db: AsyncSession,
    upload_method: str = "Bulk_Upload",
) -> list[dict]:
    """Attach up to 20 documents to a case as one bulk upload."""
    if len(files) > MAX_CASE_UPLOAD_FILES:
        _api_error("VALIDATION_ERROR", "A maximum of 20 files can be uploaded at once.", "files")
    created = []
    for item in files:
        created.append(
            await attach_document(
                case_id=case_id,
                file_name=item["file_name"],
                document_type=item["document_type"],
                file_bytes=item["file_bytes"],
                user_id=user_id,
                db=db,
                mime_type=item.get("mime_type"),
                upload_method=upload_method,
            )
        )
    return created


async def list_case_documents(case_id: str, db: AsyncSession) -> list:
    """Return all documents for a given case."""
    result = await db.execute(
        select(CaseDocument).where(CaseDocument.case_id == case_id)
    )
    return [_to_dict(d, key_map=_DOCUMENT_KEY_MAP, exclude=_DOCUMENT_EXCLUDE_KEYS) for d in result.scalars().all()]


async def get_case_document(case_id: str, doc_id: str, db: AsyncSession) -> Optional[dict]:
    """Return a single document if it belongs to *case_id*."""
    result = await db.execute(
        select(CaseDocument).where(
            CaseDocument.id == doc_id,
            CaseDocument.case_id == case_id,
        )
    )
    doc = result.scalar_one_or_none()
    if doc is None:
        return None
    data = _to_dict(doc, key_map=_DOCUMENT_KEY_MAP, exclude=_DOCUMENT_EXCLUDE_KEYS)
    parse_record_id = await _find_parse_record_id_for_case_document(doc, db)
    if parse_record_id:
        data["parse_record_id"] = parse_record_id
    return data


async def _find_parse_record_id_for_case_document(doc: CaseDocument, db: AsyncSession) -> str | None:
    """Best-effort bridge from a case document to its parse-history record."""
    candidates = []
    if doc.sha256_hash:
        candidates.append(
            sa.select(parse_records.c.id)
            .where(
                parse_records.c.case_id == doc.case_id,
                parse_records.c.file_sha256 == doc.sha256_hash,
            )
            .order_by(parse_records.c.created_at.desc())
            .limit(1)
        )
    candidates.append(
        sa.select(parse_records.c.id)
        .where(
            parse_records.c.case_id == doc.case_id,
            parse_records.c.file_name == doc.file_name,
        )
        .order_by(parse_records.c.created_at.desc())
        .limit(1)
    )
    for query in candidates:
        result = await db.execute(query)
        value = result.scalar_one_or_none()
        if value:
            return str(value)
    return None


async def list_document_versions(case_id: str, doc_id: str, db: AsyncSession) -> list[dict]:
    """Return all versions for the same case document file name."""
    doc = await db.get(CaseDocument, doc_id)
    if doc is None or doc.case_id != case_id:
        _api_error("NOT_FOUND", f"Document '{doc_id}' not found.")
    result = await db.execute(
        select(CaseDocument)
        .where(CaseDocument.case_id == case_id, CaseDocument.file_name == doc.file_name)
        .order_by(CaseDocument.version.desc())
    )
    return [_to_dict(d, key_map=_DOCUMENT_KEY_MAP, exclude=_DOCUMENT_EXCLUDE_KEYS) for d in result.scalars().all()]


async def diff_document_versions(case_id: str, left_doc_id: str, right_doc_id: str, db: AsyncSession) -> dict:
    """Return a unified diff between two stored text-like document versions."""
    import difflib

    left = await db.get(CaseDocument, left_doc_id)
    right = await db.get(CaseDocument, right_doc_id)
    if left is None or right is None or left.case_id != case_id or right.case_id != case_id:
        _api_error("NOT_FOUND", "One or both document versions were not found.")

    left_bytes = await read_case_document_bytes(left)
    right_bytes = await read_case_document_bytes(right)
    left_text = (left_bytes or b"").decode("utf-8", errors="replace").splitlines()
    right_text = (right_bytes or b"").decode("utf-8", errors="replace").splitlines()
    diff = "\n".join(
        difflib.unified_diff(
            left_text,
            right_text,
            fromfile=f"{left.file_name} v{left.version}",
            tofile=f"{right.file_name} v{right.version}",
            lineterm="",
        )
    )
    return {
        "left_document_id": left_doc_id,
        "right_document_id": right_doc_id,
        "left_version": left.version,
        "right_version": right.version,
        "diff": diff,
    }


async def purge_old_document_versions(case_id: str, user_id: str, db: AsyncSession) -> dict:
    """Delete superseded document versions (is_latest_version=False) for a case.

    Removes both the database records and their associated storage blobs.
    Returns a summary of what was purged.
    """
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    result = await db.execute(
        select(CaseDocument).where(
            CaseDocument.case_id == case_id,
            CaseDocument.is_latest_version == False,  # noqa: E712
        )
    )
    old_docs = result.scalars().all()

    if not old_docs:
        return {"purged_count": 0, "purged_documents": [], "freed_bytes": 0}

    purged: list[dict] = []
    freed_bytes = 0

    for doc in old_docs:
        uri = doc.file_storage_uri
        size = doc.file_size_bytes or 0
        purged.append({
            "id": doc.id,
            "file_name": doc.file_name,
            "version": doc.version,
            "size_bytes": size,
        })
        freed_bytes += size

        # Delete from object storage
        if uri:
            try:
                await delete_object(uri)
            except Exception:
                pass  # best-effort; record is still removed from DB

        await db.delete(doc)

    await db.flush()

    await add_activity(
        case_id, "Documents_Purged", user_id,
        f"Purged {len(purged)} old document version(s), freeing {freed_bytes:,} bytes.",
        db=db,
    )

    return {
        "purged_count": len(purged),
        "purged_documents": purged,
        "freed_bytes": freed_bytes,
    }


# ---------------------------------------------------------------------------
# Timeline / activity
# ---------------------------------------------------------------------------

async def add_activity(
    case_id: str,
    activity_type: str,
    user_id: str,
    description: str,
    entity_type: Optional[str] = None,
    entity_id: Optional[str] = None,
    *,
    db: AsyncSession,
) -> dict:
    """Append an activity entry to the case timeline."""
    activity = CaseActivity(
        case_id=case_id,
        activity_type=activity_type,
        user_id=user_id,
        description=description,
        entity_type=entity_type,
        entity_id=entity_id,
    )
    db.add(activity)
    await db.flush()
    return _to_dict(activity)


async def get_timeline(case_id: str, sort_asc: bool = False, *, db: AsyncSession) -> list:
    """Return the full activity timeline for a case (default: newest first)."""
    order = CaseActivity.created_at.asc() if sort_asc else CaseActivity.created_at.desc()
    result = await db.execute(
        select(CaseActivity)
        .where(CaseActivity.case_id == case_id)
        .order_by(order)
    )
    return [_to_dict(a) for a in result.scalars().all()]


# ---------------------------------------------------------------------------
# Action tracker
# ---------------------------------------------------------------------------

async def create_task(case_id: str, data: dict, user_id: str, db: AsyncSession) -> dict:
    """Create an action-tracker task linked to a case."""
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    due_date_raw = data.get("due_date")
    if isinstance(due_date_raw, str):
        due_date_val = date.fromisoformat(due_date_raw)
    elif isinstance(due_date_raw, date):
        due_date_val = due_date_raw
    else:
        due_date_val = None

    task = ActionTrackerTask(
        case_id=case_id,
        task_name=data.get("task_name", ""),
        due_date=due_date_val,
        priority=_normalise_priority(data.get("priority", "Medium")),
        status="Pending",
        source=data.get("source", "Manual"),
        created_by=user_id,
    )
    db.add(task)
    await db.flush()

    await add_activity(
        case_id, "Task_Created", user_id,
        f"Task '{task.task_name}' created.",
        entity_type="task", entity_id=task.id,
        db=db,
    )
    return _task_to_dict(task)


async def list_tasks(case_id: str, db: AsyncSession) -> list:
    """Return tasks for a case, sorted by due_date ascending."""
    result = await db.execute(
        select(ActionTrackerTask)
        .where(ActionTrackerTask.case_id == case_id)
        .order_by(ActionTrackerTask.due_date.asc().nullslast())
    )
    return [_task_to_dict(t) for t in result.scalars().all()]


async def update_task(case_id: str, task_id: str, data: dict, user_id: str, db: AsyncSession) -> dict:
    """Update an existing action-tracker task."""
    result = await db.execute(
        select(ActionTrackerTask).where(
            ActionTrackerTask.id == task_id,
            ActionTrackerTask.case_id == case_id,
        )
    )
    task = result.scalar_one_or_none()
    if task is None:
        _api_error("NOT_FOUND", f"Task '{task_id}' not found for case '{case_id}'.")

    mutable = ("task_name", "due_date", "priority", "status", "snoozed_until")
    for key in mutable:
        if key in data:
            val = data[key]
            # Convert string dates to proper types
            if key == "due_date" and isinstance(val, str):
                val = date.fromisoformat(val)
            elif key == "snoozed_until" and isinstance(val, str):
                val = datetime.fromisoformat(val)
            elif key == "status" and isinstance(val, str):
                val = _normalise_task_status(val)
                if val == "Completed":
                    task.completed_at = datetime.now(timezone.utc)
            elif key == "priority" and isinstance(val, str):
                val = _normalise_priority(val)
            setattr(task, key, val)

    task.updated_at = datetime.now(timezone.utc)
    await db.flush()

    await add_activity(
        case_id, "Task_Updated", user_id,
        f"Task '{task.task_name}' updated.",
        entity_type="task", entity_id=task_id,
        db=db,
    )
    return _task_to_dict(task)


async def trigger_due_task_reminders(db: AsyncSession, today: Optional[date] = None) -> list[dict]:
    """Create in-app reminders for tasks due in 3 days, 1 day, or today."""
    today = today or datetime.now(timezone.utc).date()
    result = await db.execute(select(ActionTrackerTask))
    created: list[dict] = []
    for task in result.scalars().all():
        task_data = _task_to_dict(task, today=today)
        if not task_data["reminder_due"] or not task.case_id:
            continue
        case = await db.get(Case, task.case_id)
        if case is None or not case.io_id:
            continue
        due_date = task.due_date
        days_left = (due_date - today).days if isinstance(due_date, date) else 0
        when = "today" if days_left == 0 else f"in {days_left} day{'s' if days_left != 1 else ''}"
        created.append(
            await create_notification(
                case.io_id,
                "reminder",
                f"Task '{task.task_name}' is due {when}.",
                entity_type="task",
                entity_id=task.id,
                db=db,
            )
        )
    return created


async def auto_populate_statutory_deadlines(case_id: str, offence_type: str, db: AsyncSession) -> list:
    """Create standard statutory-deadline tasks for a case.

    Deadlines are computed relative to the case creation date.
    """
    case = await db.get(Case, case_id)
    if case is None:
        _api_error("NOT_FOUND", f"Case '{case_id}' not found.")

    # Prefer FIR registration date; fall back to case creation date
    raw_date = case.date_of_registration or case.created_at
    base_date = raw_date if isinstance(raw_date, datetime) else datetime.fromisoformat(raw_date)
    templates = [
        {"task_name": "Submit progress report", "days": 30, "priority": "Medium"},
        {"task_name": "Serve summons to accused", "days": 60, "priority": "High"},
        {"task_name": "Witness examination completion", "days": 60, "priority": "Medium"},
        {"task_name": "File charge sheet", "days": 90, "priority": "High"},
        {"task_name": "FSL report follow-up", "days": 45, "priority": "Medium"},
    ]

    io_id = case.io_id if isinstance(case.io_id, str) else str(case.io_id)

    created: list = []
    for tpl in templates:
        due = (base_date + timedelta(days=tpl["days"])).date().isoformat()
        task = await create_task(
            case_id,
            {
                "task_name": tpl["task_name"],
                "due_date": due,
                "priority": tpl["priority"],
                "source": "Statutory",
            },
            io_id,
            db,
        )
        created.append(task)

    return created


# ---------------------------------------------------------------------------
# Notifications
# ---------------------------------------------------------------------------

async def create_notification(
    user_id: str,
    type: str,
    message: str,
    entity_type: Optional[str] = None,
    entity_id: Optional[str] = None,
    *,
    db: AsyncSession,
) -> dict:
    """Create an in-app notification for a user."""
    note = Notification(
        user_id=user_id,
        type=type,
        message=message,
        entity_type=entity_type,
        entity_id=entity_id,
        is_read=False,
    )
    db.add(note)
    await db.flush()
    return _to_dict(note)


async def list_notifications(user_id: str, unread_only: bool = True, *, db: AsyncSession) -> list:
    """Return notifications for a user, newest first."""
    stmt = select(Notification).where(Notification.user_id == user_id)
    if unread_only:
        stmt = stmt.where(Notification.is_read == False)  # noqa: E712
    stmt = stmt.order_by(Notification.created_at.desc())
    result = await db.execute(stmt)
    return [_to_dict(n) for n in result.scalars().all()]


async def mark_notification_read(notification_id: str, db: AsyncSession) -> dict:
    """Mark a single notification as read."""
    note = await db.get(Notification, notification_id)
    if note is None:
        _api_error("NOT_FOUND", f"Notification '{notification_id}' not found.")
    note.is_read = True
    await db.flush()
    return _to_dict(note)


# ---------------------------------------------------------------------------
# Seed data — police stations & offence types
# ---------------------------------------------------------------------------

async def list_police_stations_data(db: AsyncSession) -> list:
    """Return all active police stations."""
    result = await db.execute(
        select(PoliceStation).where(PoliceStation.is_active == True)  # noqa: E712
    )
    return [_to_dict(ps, key_map=_POLICE_STATION_KEY_MAP) for ps in result.scalars().all()]


async def list_offence_types_data(db: AsyncSession) -> list:
    """Return all active offence types."""
    result = await db.execute(
        select(OffenceType).where(OffenceType.is_active == True)  # noqa: E712
    )
    return [_to_dict(o) for o in result.scalars().all()]


async def search_offence_types_data(query: str, db: AsyncSession) -> list:
    """Search offence types by name (case-insensitive substring match)."""
    result = await db.execute(
        select(OffenceType)
        .where(OffenceType.is_active == True)  # noqa: E712
        .where(OffenceType.name.ilike(f"%{query}%"))
    )
    return [_to_dict(o) for o in result.scalars().all()]


async def seed_police_stations(db: AsyncSession) -> None:
    """Populate sample Hyderabad police stations (idempotent)."""
    stations = [
        {"id": "ps-001", "name": "Banjara Hills PS", "station_code": "BH", "district": "Hyderabad"},
        {"id": "ps-002", "name": "Begumpet PS", "station_code": "BG", "district": "Hyderabad"},
        {"id": "ps-003", "name": "Jubilee Hills PS", "station_code": "JH", "district": "Hyderabad"},
        {"id": "ps-004", "name": "Madhapur PS", "station_code": "MP", "district": "Cyberabad"},
        {"id": "ps-005", "name": "Saifabad PS", "station_code": "SF", "district": "Hyderabad"},
    ]
    for s in stations:
        existing = await db.get(PoliceStation, s["id"])
        if existing is None:
            db.add(PoliceStation(**s))
    await db.flush()


async def seed_offence_types(db: AsyncSession) -> None:
    """Populate common offence types with BNS/IPC sections (idempotent)."""
    offences = [
        {"id": "off-001", "name": "Theft", "bns_section": "303", "ipc_section": "379"},
        {"id": "off-002", "name": "Robbery", "bns_section": "309", "ipc_section": "392"},
        {"id": "off-003", "name": "Cheating", "bns_section": "318", "ipc_section": "420"},
        {"id": "off-004", "name": "Criminal Breach of Trust", "bns_section": "316", "ipc_section": "406"},
        {"id": "off-005", "name": "Assault", "bns_section": "115", "ipc_section": "323"},
        {"id": "off-006", "name": "Murder", "bns_section": "101", "ipc_section": "302"},
        {"id": "off-007", "name": "Kidnapping", "bns_section": "137", "ipc_section": "363"},
        {"id": "off-008", "name": "Dowry Death", "bns_section": "80", "ipc_section": "304B"},
        {"id": "off-009", "name": "Cyber Crime", "bns_section": "318", "ipc_section": "66 IT Act"},
        {"id": "off-010", "name": "Criminal Intimidation", "bns_section": "351", "ipc_section": "506"},
    ]
    for o in offences:
        existing = await db.get(OffenceType, o["id"])
        if existing is None:
            db.add(OffenceType(**o))
    await db.flush()
